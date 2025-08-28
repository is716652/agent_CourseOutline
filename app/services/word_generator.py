#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档生成器
将生成的教学大纲内容回填到Word模板
"""

import os
import re
from docx import Document
from loguru import logger


def generate_word_document(outline_data, template_path, output_path):
    """
    生成Word文档
    
    Args:
        outline_data: 教学大纲数据字典
        template_path: Word模板文件路径
        output_path: 输出Word文件路径
    
    Returns:
        str: 生成的文件路径
    """
    
    try:
        # 读取Word模板
        doc = Document(template_path)
        logger.info(f"加载Word模板: {template_path}")
        
        # 替换段落中的变量
        for paragraph in doc.paragraphs:
            replace_variables_in_paragraph(paragraph, outline_data)
        
        # 替换表格中的变量
        for table in doc.tables:
            replace_variables_in_table(table, outline_data)
        
        # 保存文档
        doc.save(output_path)
        logger.info(f"Word文档已生成: {output_path}")
        
        return output_path
        
    except Exception as e:
        logger.error(f"生成Word文档失败: {e}")
        raise


def replace_variables_in_paragraph(paragraph, data):
    """替换段落中的模板变量，保持原有格式"""
    
    # 获取段落的完整文本
    full_text = paragraph.text
    
    # 查找所有模板变量
    variables = re.findall(r'\{\{([^}]+)\}\}', full_text)
    
    if not variables:
        return
    
    # 使用高级方法处理跨多个 runs 的模板变量
    replace_variables_advanced(paragraph, data)


def replace_variables_advanced(paragraph, data):
    """高级模板变量替换，处理跨多个 runs 的情况"""
    
    # 获取段落中所有的 runs 信息
    runs_info = []
    for run in paragraph.runs:
        runs_info.append({
            'text': run.text,
            'font_name': run.font.name,
            'font_size': run.font.size,
            'font_bold': run.font.bold,
            'font_italic': run.font.italic,
            'font_underline': run.font.underline,
            'font_color': run.font.color.rgb if run.font.color else None
        })
    
    # 合并所有 runs 的文本
    full_text = ''.join([info['text'] for info in runs_info])
    
    # 查找并替换模板变量
    variables = re.findall(r'\{\{([^}]+)\}\}', full_text)
    if not variables:
        return
    
    # 替换模板变量
    new_text = full_text
    for var_name in variables:
        var_value = data.get(var_name, generate_default_value(var_name))
        new_text = new_text.replace(f"{{{{{var_name}}}}}", str(var_value))
    
    # 如果文本没有变化，不做处理
    if new_text == full_text:
        return
    
    # 清空段落
    paragraph.clear()
    
    # 尝试保持最后一个 run 的格式（大多数情况下模板变量会在最后）
    if runs_info:
        last_run_info = runs_info[-1]
        new_run = paragraph.add_run(new_text)
        
        # 应用最后一个 run 的格式
        if last_run_info['font_name']:
            new_run.font.name = last_run_info['font_name']
        if last_run_info['font_size']:
            new_run.font.size = last_run_info['font_size']
        if last_run_info['font_bold'] is not None:
            new_run.font.bold = last_run_info['font_bold']
        if last_run_info['font_italic'] is not None:
            new_run.font.italic = last_run_info['font_italic']
        if last_run_info['font_underline'] is not None:
            new_run.font.underline = last_run_info['font_underline']
        if last_run_info['font_color']:
            new_run.font.color.rgb = last_run_info['font_color']
    else:
        # 如果没有原有 runs，只是添加简单文本
        paragraph.add_run(new_text)


def generate_default_value(var_name):
    """为缺失的变量生成合理的默认值"""
    
    # 根据变量名称生成对应的默认内容
    if '教学模块' in var_name:
        module_num = re.search(r'(\d+)', var_name)
        if module_num:
            num = module_num.group(1)
            module_names = {
                '3': '实践应用与操作',
                '4': '项目开发与设计', 
                '5': '案例分析与研究',
                '6': '综合实训与考核',
                '7': '前沿技术与发展',
                '8': '总结提升与拓展'
            }
            return module_names.get(num, f'教学模块{num}')
    
    elif '教学内容及重点、难点' in var_name:
        module_num = re.search(r'(\d+)', var_name)
        if module_num:
            num = module_num.group(1)
            content_templates = {
                '3': '重点：实践操作技能、应用方法；难点：理论与实践的结合应用',
                '4': '重点：项目设计方法、开发流程；难点：综合运用知识进行项目设计',
                '5': '重点：典型案例分析、问题研究方法；难点：批判性思维和分析能力培养',
                '6': '重点：综合实训项目、技能考核标准；难点：综合技能的整合应用',
                '7': '重点：前沿技术发展、新技术应用；难点：新技术的理解和应用前景分析',
                '8': '重点：知识体系总结、能力提升路径；难点：知识整合和能力提升规划'
            }
            return content_templates.get(num, f'教学内容及重点难点{num}的具体描述')
    
    elif '职业技能要求' in var_name:
        module_num = re.search(r'(\d+)', var_name)
        if module_num:
            num = module_num.group(1)
            skill_templates = {
                '3': '具备独立操作和应用能力',
                '4': '具备项目设计和开发能力',
                '5': '具备案例分析和研究能力',
                '6': '达到综合应用技能标准',
                '7': '了解行业发展和新技术趋势',
                '8': '形成完整的知识技能体系'
            }
            return skill_templates.get(num, f'掌握模块{num}相关技能')
    
    elif '课时' in var_name:
        module_num = re.search(r'(\d+)', var_name)
        if module_num:
            return '8学时'
    
    elif '教学方法建议' in var_name:
        module_num = re.search(r'(\d+)', var_name)
        if module_num:
            num = module_num.group(1)
            method_templates = {
                '3': '实践操作结合指导训练',
                '4': '项目驱动教学法',
                '5': '案例教学法结合小组讨论',
                '6': '综合实训结合技能考核',
                '7': '专题讲座结合技术调研',
                '8': '总结交流结合学习规划'
            }
            return method_templates.get(num, '理论讲授结合实践操作')
    
    # 其他变量的默认值
    return f'[请填写{var_name}]'


def replace_variables_in_table(table, data):
    """替换表格中的模板变量，保持表格格式"""
    
    for row in table.rows:
        for cell in row.cells:
            # 保存原有的单元格格式
            for paragraph in cell.paragraphs:
                replace_variables_in_paragraph(paragraph, data)
            
            # 保持单元格的其他属性（如背景颜色、边框等）
            # 这些属性在替换文本时不会被改变


def create_word_from_outline(outline_data, course_name, output_dir):
    """
    从教学大纲数据创建Word文档
    
    Args:
        outline_data: 教学大纲数据
        course_name: 课程名称
        output_dir: 输出目录
    
    Returns:
        str: 生成的Word文件路径
    """
    
    # 构建文件路径
    project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '../..'))
    template_path = os.path.join(project_root, 'templates', '教学大纲-模板.docx')
    
    # 清理课程名称，用于文件名
    safe_course_name = clean_filename(course_name)
    output_filename = f"教学大纲-{safe_course_name}.docx"
    output_path = os.path.join(output_dir, output_filename)
    
    # 确保输出目录存在
    os.makedirs(output_dir, exist_ok=True)
    
    # 生成Word文档
    return generate_word_document(outline_data, template_path, output_path)


def clean_filename(filename):
    """清理文件名，移除不合法字符"""
    
    # 移除不合法字符
    illegal_chars = '<>:"/\\|?*'
    for char in illegal_chars:
        filename = filename.replace(char, '')
    
    # 限制长度
    return filename[:50] if len(filename) > 50 else filename


if __name__ == "__main__":
    # 测试代码
    test_data = {
        '课程名称': '计算机网络基础',
        '编写日期': '2024年3月',
        '课程定位': '本课程是计算机类专业的核心课程',
        '知识目标': '掌握计算机网络的基本概念和原理',
        '技能目标': '具备网络配置和管理能力',
        '素质目标': '培养网络安全意识和团队协作精神'
    }
    
    output_dir = "output"
    result = create_word_from_outline(test_data, test_data['课程名称'], output_dir)
    print(f"测试生成完成: {result}")